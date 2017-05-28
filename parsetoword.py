from datetime import datetime

from docx import Document
from docx.shared import Inches

import urllib2, StringIO

def convert_to_word(entry, output):
    document = Document()

    document.add_heading(entry.title, 0)
    document.add_heading(datetime.strftime(entry.date, '%Y-%m-%d %H:%M:%S'), level=1)
    document.add_heading(entry.link, level=1)

    for content in entry.contents:
        content.change_document(document)
    document.save("%s/%s.docx" % (output, datetime.strftime(entry.date, '%Y-%m-%d %H%M%S')))

class AmebloEntry(object):
    title = ""
    link = ""
    date = datetime.now()
    contents = []

class AmebloContent(object):
    def get_content(self):
        pass

    def change_document(self, document):
        pass

class AmebloImage(AmebloContent):
    def __init__(self, img_url):
        self.img_url = img_url

    def get_content(self):
        image_from_url = urllib2.urlopen(self.img_url)
        io_url = StringIO.StringIO()
        io_url.write(image_from_url.read())
        io_url.seek(0)
        return io_url

    def change_document(self, document):
        try:
            document.add_picture(self.get_content())
        except:
            pass

class AmebloText(AmebloContent):
    def __init__(self, text):
        self.text = text

    def get_content(self):
        return self.text

    def change_document(self, document):
        document.add_paragraph(self.text)